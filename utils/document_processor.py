import os
import PyPDF2
import docx
from typing import List

try:
    # LangChain >= 0.2 style
    from langchain_core.documents import Document
except ImportError:
    # Legacy path
    from langchain.schema import Document

class DocumentProcessor:
    def __init__(self, window_pages: int = 2, overlap_pages: int = 1):
        self.window_pages = max(window_pages, 1)
        self.overlap_pages = max(min(overlap_pages, self.window_pages - 1), 0)

    def extract_pages_from_pdf(self, file_path: str) -> List[str]:
        """Extract text page-by-page from PDF."""
        pages = []
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    pages.append(page.extract_text() or "")
        except Exception as e:
            print(f"Error reading PDF {file_path}: {str(e)}")
        return pages

    def extract_single_page_from_docx(self, file_path: str) -> List[str]:
        """Treat entire DOCX as a single 'page'."""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {str(e)}")
        return [text]

    def extract_single_page_from_txt(self, file_path: str) -> List[str]:
        """Treat entire TXT as a single 'page'."""
        text = ""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()
        except Exception as e:
            print(f"Error reading TXT {file_path}: {str(e)}")
        return [text]

    def page_windows(self, pages: List[str]) -> List[Document]:
        """Create overlapping page windows (e.g., 1-2, 2-3, 3-4...)."""
        if not pages:
            return []
        step = max(self.window_pages - self.overlap_pages, 1)
        docs: List[Document] = []
        total = len(pages)
        for start in range(0, total, step):
            end = min(start + self.window_pages, total)
            window_text = "\n".join(pages[start:end]).strip()
            if not window_text:
                continue
            if end == start + 1:
                page_range = f"page {start+1}"
            else:
                page_range = f"pages {start+1}-{end}"
            docs.append(
                Document(
                    page_content=window_text,
                    metadata={
                        "page_range": page_range,
                    },
                )
            )
            if end == total:
                break
        return docs

    def process_document(self, file_path: str) -> List[Document]:
        """Process a document and return page-windowed documents."""
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == ".pdf":
            pages = self.extract_pages_from_pdf(file_path)
        elif file_extension == ".docx":
            pages = self.extract_single_page_from_docx(file_path)
        elif file_extension == ".txt":
            pages = self.extract_single_page_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

        if not any(p.strip() for p in pages):
            raise ValueError(f"No text extracted from {file_path}")

        docs = self.page_windows(pages)
        filename = os.path.basename(file_path)
        for d in docs:
            d.metadata["source"] = file_path
            d.metadata["filename"] = filename
        return docs

    def process_multiple_documents(self, file_paths: List[str]) -> List[Document]:
        """Process multiple documents and return all chunks."""
        all_chunks = []
        for file_path in file_paths:
            try:
                chunks = self.process_document(file_path)
                all_chunks.extend(chunks)
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
                continue
        return all_chunks
